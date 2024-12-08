from faker import Faker
import random
import os
import pathlib
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_academic_report(
    faker, save_dir=r"files\synthesised_word_documents", mask_type=0
):

    report = docx.Document()

    # report format -> title -> contents -> abstract -> topics -> bibliography

    report_title = faker.catch_phrase()
    report_author_1 = faker.name()
    report_author_2 = faker.name()
    report_author_3 = faker.name()

    h1 = report.add_heading(report_title)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = report.add_heading(
        ", ".join([report_author_1, report_author_2, report_author_3]), level=2
    )
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.add_page_break()
    report.add_section()

    report.add_heading("Abstract")
    report.add_paragraph(faker.text())

    report.add_page_break()
    report.add_section()

    report.add_heading("Table of Contents")
    table = report.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Topic."
    hdr_cells[1].text = "Page No."

    num_topics = random.choice(range(5, 10))
    for i in range(0, num_topics):
        row_cells = table.add_row().cells
        row_cells[0].text = faker.sentence()
        row_cells[1].text = str(i + 1)

    for cell in table.column_cells(0)[1:]:
        report.add_page_break()
        report.add_section()
        report.add_heading(cell.text)
        report.add_paragraph(faker.paragraph(nb_sentences=random.choice(range(10, 50))))

    report.add_page_break()
    report.add_section()

    report.add_heading("Bibliography")
    refs = report.add_table(rows=1, cols=2)
    hdr_cells = refs.rows[0].cells
    hdr_cells[0].text = "No."
    hdr_cells[1].text = "Description"

    num_topics = random.choice(range(2, 7))
    for i in range(0, num_topics):
        row_cells = refs.add_row().cells
        row_cells[0].text = "[" + str(i + 1) + "]"
        row_cells[1].text = faker.sentence()

    if mask_type == 0:
        prefix = "academic_report"
    else:
        prefix = "document"

    report.save(rf"{save_dir}\{prefix}_{faker.random_number()}.docx")


def generate_business_report(
    faker, save_dir=r"files\synthesised_word_documents", mask_type=0
):

    report = docx.Document()

    # report format -> title -> exec. summary -> contents -> topics -> refs -> appendix

    report_type = random.choice(["Annual", "Monthly"])
    if report_type == "Annual":
        report_period = faker.year()
    if report_type == "Monthly":
        report_period = faker.month_name() + " " + faker.year()

    h1 = report.add_heading(f"{report_type} Report: {report_period}")
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.add_page_break()
    report.add_section()

    report.add_heading("Executive Summary")
    report.add_paragraph(faker.text())

    report.add_page_break()
    report.add_section()

    report.add_heading("Table of Contents")
    table = report.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Topic."
    hdr_cells[1].text = "Page No."

    num_topics = random.choice(range(3, 6))
    for i in range(0, num_topics):
        if i == 0:
            row_cells = table.add_row().cells
            row_cells[0].text = "Introduction"
            row_cells[1].text = str(i + 1)
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = faker.sentence()
            row_cells[1].text = str(i + 1)

    for cell in table.column_cells(0)[1:]:
        report.add_page_break()
        report.add_section()
        report.add_heading(cell.text)
        report.add_paragraph(faker.paragraph(nb_sentences=random.choice(range(25, 50))))

    report.add_page_break()
    report.add_section()

    report.add_heading("References")
    refs = report.add_table(rows=1, cols=2)
    hdr_cells = refs.rows[0].cells
    hdr_cells[0].text = "Ref. No."
    hdr_cells[1].text = "Description"

    num_topics = random.choice(range(2, 7))
    for i in range(0, num_topics):
        row_cells = refs.add_row().cells
        row_cells[0].text = "[" + str(i + 1) + "]"
        row_cells[1].text = faker.sentence()

    report.add_page_break()
    report.add_section()

    report.add_page_break()
    report.add_section()
    report.add_heading("Appendix")
    report.add_paragraph(faker.paragraph(nb_sentences=random.choice(range(25, 50))))

    if mask_type == 0:
        prefix = "business_report"
    else:
        prefix = "document"

    report.save(rf"{save_dir}\{prefix}_{faker.random_number()}.docx")


def generate_formal_letter(
    faker, save_dir=r"files\synthesised_word_documents", mask_type=0
):

    letter = docx.Document()

    sender = faker.name()
    sender_address = faker.address()

    receiver = faker.name()
    receiver_address = faker.address()

    letter_date = faker.date_time_this_year().strftime("%A, %d %b %Y")

    subject = faker.catch_phrase()

    p1 = letter.add_paragraph(
        f"""{sender}
{sender_address}

{letter_date}
"""
    )

    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    letter.add_paragraph(
        f"""{receiver}
{receiver_address}
"""
    )

    letter_contents = f"""Dear {receiver},

    Subject: {subject}

{faker.paragraph(nb_sentences=4)}

{faker.paragraph(nb_sentences=8)}

{faker.paragraph(nb_sentences=3)}

Yours Sincerely,
{sender}
"""

    letter.add_paragraph(letter_contents)

    if mask_type == 0:
        prefix = "formal_letter"
    else:
        prefix = "document"

    letter.save(rf"{save_dir}\{prefix}_{faker.random_number()}.docx")


def create_reports(
    faker, doc_type="", save_dir=r"files\synthesised_word_documents", num=1, mask_type=0
):
    if doc_type == "academic_report":
        for i in range(0, num):
            generate_academic_report(faker, save_dir=save_dir, mask_type=mask_type)
    if doc_type == "business_report":
        for i in range(0, num):
            generate_business_report(faker, save_dir=save_dir, mask_type=mask_type)
    if doc_type == "formal_letter":
        for i in range(0, num):
            generate_formal_letter(faker, save_dir=save_dir, mask_type=mask_type)


if __name__ == "__main__":

    faker = Faker()

    # 1. creating files and saving to files/x/

    # create_reports(faker, doc_type="academic_report", num=2)
    # create_reports(faker, doc_type="business_report", num=2)
    # create_reports(faker, doc_type="formal_letter", num=2)
    # create_reports(faker, doc_type="academic_report", num=2, mask_type=1)
    # create_reports(faker, doc_type="business_report", num=2, mask_type=1)
    # create_reports(faker, doc_type="formal_letter", num=2, mask_type=1)

    # create_reports(faker, doc_type="formal_letter", save_dir="files",num=2, mask_type=1)

    # 2. creating training data and saving to files/synthesised_training_data/x/

    # acad_rep_output_dir = r"files\synthesised_training_data\acad_rep_docx"
    # bus_rep_output_dir = r"files\synthesised_training_data\bus_rep_docx"
    # letter_output_dir = r"files\synthesised_training_data\fml_let_docx"

    # pathlib.Path(acad_rep_output_dir).mkdir(parents=True, exist_ok=True)
    # pathlib.Path(bus_rep_output_dir).mkdir(parents=True, exist_ok=True)
    # pathlib.Path(letter_output_dir).mkdir(parents=True, exist_ok=True)

    # num_files = 250

    # create_reports(
    #     faker,
    #     doc_type="academic_report",
    #     save_dir=acad_rep_output_dir,
    #     num=num_files,
    #     mask_type=1,
    # )
    # create_reports(
    #     faker,
    #     doc_type="business_report",
    #     save_dir=bus_rep_output_dir,
    #     num=num_files,
    #     mask_type=1,
    # )
    # create_reports(
    #     faker,
    #     doc_type="formal_letter",
    #     save_dir=letter_output_dir,
    #     num=num_files,
    #     mask_type=1,
    # )

    # 3. creating testing data to be used for predicton and saving to files/synthesised_test_data/x

    acad_rep_output_dir = r"files\synthesised_test_data\acad_rep_test_docx"
    bus_rep_output_dir = r"files\synthesised_test_data\bus_rep_test_docx"
    letter_output_dir = r"files\synthesised_test_data\fml_let_test_docx"

    pathlib.Path(acad_rep_output_dir).mkdir(parents=True, exist_ok=True)
    pathlib.Path(bus_rep_output_dir).mkdir(parents=True, exist_ok=True)
    pathlib.Path(letter_output_dir).mkdir(parents=True, exist_ok=True)

    num_files = 10

    create_reports(
        faker,
        doc_type="academic_report",
        save_dir=acad_rep_output_dir,
        num=num_files,
        mask_type=1,
    )
    create_reports(
        faker,
        doc_type="business_report",
        save_dir=bus_rep_output_dir,
        num=num_files,
        mask_type=1,
    )
    create_reports(
        faker,
        doc_type="formal_letter",
        save_dir=letter_output_dir,
        num=num_files,
        mask_type=1,
    )
